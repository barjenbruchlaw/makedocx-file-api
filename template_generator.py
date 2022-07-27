from docx import Document
from datetime import datetime

#Creates new dictionary to which to add template parameters
template_dict = dict()

#Create timedate string for default parameter on output file name
time_now = datetime.now()
time_stamp = time_now.strftime('%Y-%m-%d_%H%M%S')

#Main function start
def template_generator(docx_template_file,docx_template_name,docx_output_filename='New File '+time_stamp):

#   Adds the provided template path and output file name to the template parameter dict
    template_dict.update({
        'template_path_field': docx_template_file,
        'template_name_field': docx_template_name,
        'output_filename_field': docx_output_filename,
    })

#   Starts a new list to hold runs to be replaced
    replacement_paragraph_field_list = list()

#   Loads document from first parameter
    document = Document(docx_template_file)

#   Breaks document into paragraphs
    paragraphs = document.paragraphs

    for paragraph_count,paragraph in enumerate(paragraphs):

#       Breaks paragraphs into runs
        runs = paragraph.runs
        for run_count, run in enumerate(runs):

#       If replacement string marker is found, adds new dict entry to replacement field list with all of the run elements
            if "{{" in run.text:
                run_dict = dict()
                run_dict.update({'paragraph':paragraph_count,'run': run_count,'text': run.text})
                replacement_paragraph_field_list.append(run_dict)

#   Adds runs to be replaced to main dictionary
    template_dict.update({'update_paragraph_runs_field': replacement_paragraph_field_list})

#   Performs similar function for tables.  Tables are broken into rows, cells, paragraphs and runs.

    replacement_table_field_list = list()

    tables = document.tables

    for table_count, table in enumerate(tables):
        rows = table.rows
        for row_count, row in enumerate(rows):
            cells = row.cells
            for cell_count, cell in enumerate(cells):

#               The prior operation is basically repeated here with new variable names.
                cell_paragraphs = cell.paragraphs
                for cell_paragraph_count, cell_paragraph in enumerate(cell_paragraphs):
                    cell_runs = cell_paragraph.runs
                    for cell_run_count, cell_run in enumerate(cell_runs):
                        if "{{" in cell_run.text:
                            cell_run_dict = dict()

#                           The additional elements are tracked in the dictionary.
                            cell_run_dict.update({
                                'table':table_count,
                                'row':row_count,
                                'cell':cell_count,
                                'paragraph':cell_paragraph_count,
                                'run':cell_run_count,
                                'text':cell_run.text
                            })
                            replacement_table_field_list.append(cell_run_dict)
    template_dict.update({'update_table_runs_field': replacement_table_field_list})

    return template_dict
