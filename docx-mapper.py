from docx import Document

document = Document("Judgment - Missouri.docx")

# document_style_list = document.styles
#
# for style in document_style_list:
#     print(style)

# sections = document.sections
#
# for section in sections:
#     print(section)

# paragraphs = document.paragraphs
#
# for paragraph in paragraphs:
#     print(paragraph)

tables = document.tables

table_map = list()

for table_count, table in enumerate(tables):
    table_map.append(list())
    table_map_table = table_map[table_count]
    rows = table.rows
    for row_count, row in enumerate(rows):
        table_map_table.append(list())
        table_map_row = table_map_table[row_count]
        cells = row.cells
        for cell_count, cell in enumerate(cells):
            table_map_row.append(list())
            table_map_cell = table_map_row[cell_count]
            paragraphs = cell.paragraphs
            for paragraph_count, paragraph in enumerate(paragraphs):
                table_map_cell.append(list())
                table_map_cell_paragraph = table_map_cell[paragraph_count]
                runs = paragraph.runs
                for run_count, run in enumerate(runs):
                    run_element_list = ['bold', 'italic', 'underline', 'font.subscript', 'font.superscript',
                                        'font.highlight_color', 'font.strike', 'font.double_strike', 'font.emboss',
                                        'font.imprint', 'font.outline', 'font.shadow', 'font.name', 'font.size']
                    table_map_cell_paragraph.append([run_count, dict()])
                    table_map_run_entry = table_map_cell_paragraph[run_count][1]
                    table_map_run_entry.update({'text': run.text})
                    for run_element in run_element_list:
                        run_element_search = f'run.{run_element}'
                        if eval(run_element_search, {}, {'run': run}) != None:
                            table_map_run_entry.update({run_element: eval(run_element_search, {}, {'run': run})})

print(table_map)
