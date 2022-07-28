from docx import Document

document = Document('22XX-000000 - Judgment.docx')

paragraphs = document.paragraphs

for paragraph_count,paragraph in enumerate(paragraphs):
    runs = paragraph.runs
    for run_count,run in enumerate(runs):
        if run.text == '\n':
            print(f'Empty new line --> paragraph: {paragraph_count}, run: {run_count}')
        if run.text == '\r':
            print(f'Empty return --> paragraph: {paragraph_count}, run: {run_count}')
        if run.text == '':
            print(f'Empty string --> paragraph: {paragraph_count}, run: {run_count}')
        if run.text == '\s':
            print(f'Empty space --> paragraph: {paragraph_count}, run: {run_count}')

tables = document.tables
for table_count,table in enumerate(tables):
    rows = table.rows
    for row_count,row in enumerate(rows):
        cells = row.cells
        for cell_count,cell in enumerate(cells):
            paragraphs = cell.paragraphs
            for paragraph_count, paragraph in enumerate(paragraphs):
                runs = paragraph.runs
                for run_count, run in enumerate(runs):
                    if run.text == '\n':
                        print(f'Empty new line --> table: {table_count}, row: {row_count}, cell: {cell_count}, paragraph: {paragraph_count}, run: {run_count}')
                    if run.text == '\r':
                        print(f'Empty return --> table: {table_count}, row: {row_count}, cell: {cell_count}, paragraph: {paragraph_count}, run: {run_count}')
                    if run.text == '':
                        print(f'Empty string --> table: {table_count}, row: {row_count}, cell: {cell_count}, paragraph: {paragraph_count}, run: {run_count}')
                    if run.text == '\s':
                        print(f'Empty space --> table: {table_count}, row: {row_count}, cell: {cell_count}, paragraph: {paragraph_count}, run: {run_count}')

target_item = tables[0].rows[2].cells[0]

target_item_paragraphs = target_item.paragraphs

for target_paragraph_count,target_paragraph in enumerate(target_item_paragraphs):
    print(f'paragraph: {target_paragraph_count}, text: {repr(target_paragraph.text)}')

target_item_new = tables[0].rows[2].cells[0]
for target_paragraph_count,target_paragraph in enumerate(target_item_new.paragraphs):
    target_item_runs = target_paragraph.runs
    for target_run_count,target_run in enumerate(target_item_runs):
        print(f'paragraph: {target_paragraph_count}, run: {target_run_count}, text: {repr(target_run.text)}')

