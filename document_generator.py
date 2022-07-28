import json
from docx import Document
from datetime import datetime
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from database import get_templates_function
from sample_values import sample_values
from rebuild_run import rebuild_run
from find_replacement_terms import find_terms
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy import Column, Integer, String
from sqlalchemy.dialects.postgresql import JSONB


SQLALCHEMY_DATABASE_URL = 'postgresql://postgres:Barton1993!@localhost:5432/templates'
engine = create_engine(SQLALCHEMY_DATABASE_URL)
Session = sessionmaker()
Session.configure(bind=engine,autocommit=False,autoflush=False)
session = Session()

Base = declarative_base()

class Template(Base):
    __tablename__ = "template"
    template_id = Column(Integer,primary_key=True)
    template_path = Column(String)
    template_name = Column(String)
    output_filename = Column(String)
    update_paragraph_runs=Column(JSONB)
    update_table_runs=Column(JSONB)

print(get_templates_function())
print("Choose template ID number")
template_id_choice = input()

template_entry = session.query(Template).filter_by(template_id=template_id_choice).first()

update_paragraph_runs_final = template_entry.update_paragraph_runs
update_table_runs_final = template_entry.update_table_runs

term_dictionary = sample_values
time_now = datetime.now()
time_stamp = time_now.strftime('%Y-%m-%d_%H%M%S')
term_dictionary.update({"{{datetimestamp}}": time_stamp})

template_doc=template_entry.template_path

template_document = Document(template_doc)

template_paragraphs = template_document.paragraphs

for template_paragraph_count,template_paragraph in enumerate(template_paragraphs):
    template_runs = template_paragraph.runs
    for template_run_count,template_run in enumerate(template_runs):
          for update_paragraph_run in update_paragraph_runs_final:
            if update_paragraph_run["paragraph"] == template_paragraph_count and update_paragraph_run["run"] == template_run_count:
                new_run = rebuild_run(template_run,term_dictionary)
                if new_run.text == '':
                    print("Blank run!")
                    template_run.clear()
                else:
                    template_paragraph.runs[template_run_count] = new_run

template_tables = template_document.tables

for template_table_count,template_table in enumerate(template_tables):
    template_table_rows = template_table.rows
    for template_table_row_count,template_table_row in enumerate(template_table_rows):
        template_table_cells = template_table_row.cells
        for template_table_cell_count,template_table_cell in enumerate(template_table_cells):
            template_table_paragraphs = template_table_cell.paragraphs
            for template_table_paragraph_count, template_table_paragraph in enumerate(template_table_paragraphs):
                template_table_runs = template_table_paragraph.runs
                for template_table_run_count, template_table_run in enumerate(template_table_runs):
                    for update_table_run in update_table_runs_final:
                        if update_table_run["table"] == template_table_count and \
                                update_table_run["row"] == template_table_row_count and \
                                update_table_run["cell"] == template_table_cell_count and \
                                update_table_run["paragraph"] == template_table_paragraph_count and \
                                update_table_run["run"] == template_table_run_count:
                            new_run = rebuild_run(template_table_run, term_dictionary)
                            template_table_paragraph.runs[template_table_run_count] = new_run

output_filename_template=template_entry.output_filename
output_filename_entry = [output_filename_template]
output_filename_dynamic_terms = find_terms(output_filename_template)
for output_term in output_filename_dynamic_terms:
    new_dynamic_term = term_dictionary[output_term]
    new_output_filename_entry=output_filename_entry[-1].replace(output_term,new_dynamic_term)
    output_filename_entry.append(new_output_filename_entry)

save_filename=output_filename_entry[-1]

template_document.save(save_filename)
