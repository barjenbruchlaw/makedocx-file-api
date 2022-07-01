from dataclasses import replace
from datetime import datetime
from docx import Document

document = Document("Judgment - Jackson County.docx")

paragraphs = document.paragraphs

paragraph_count = len(paragraphs)

table_count = len(document.tables)

now = datetime.now()
dt_string = now.strftime("%Y-%m-%d %H:%M:%S")

static_term_dict = {
    "{{Street Address}}": "12345 Main Street",
    "{{City}}": "Kansas City",
    "{{ZIP}}": "64114",
    "{{Monthly Rent}}": "$1,000.00",
}

replacement_dict = dict()

for a in range(len(paragraphs)):
    next_line = paragraphs[a].text
    if "{{" in next_line:
        braces_count = next_line.count("{{")
        all_braces_index = list()
        search_term_index = list()
        for b in range(len(next_line)):
            c = b+2
            if next_line[b:c] == "{{":
                all_braces_index.append(b)
        for d in range(len(next_line)):
            e = d+2
            if next_line[d:e] == "}}":
                all_braces_index.append((d+2))
        all_braces_index.sort()
        i = 0
        while i < len(all_braces_index):
            next_search_term = next_line[all_braces_index[i]:all_braces_index[i+1]]
            search_term_index.append(next_search_term)
            i += 2
        replacement_dict.update(
            {a: [all_braces_index, search_term_index, next_line]})

j = 0

while j < paragraph_count:
    edit_paragraph = paragraphs[j].text
    if j in replacement_dict:
        for each_term in replacement_dict[j][1]:
            replaced_term = each_term
            if replaced_term in static_term_dict:
                replacing_term = static_term_dict[replaced_term]
                new_paragraph=edit_paragraph.replace(replaced_term, replacing_term)
            else:
                new_paragraph=edit_paragraph.replace(replaced_term, "")
            edit_paragraph=new_paragraph
        paragraphs[j].clear()
        paragraphs[j].add_run(text=new_paragraph)
    j+=1

key_list = list()

for f in replacement_dict:
    key_list.append(f)

key_list.sort()

last_key = key_list[-1]

table_matrix_dict = dict()

if table_count > 0:
    table_matrix = []
    for x, table in enumerate(document.tables):
        for row in table.rows:
            table_row = []
            for cell in row.cells:
                table_row.append(cell.text)
            table_matrix.append(table_row)
    table_matrix_dict[x] = table_matrix

print("paragraph count = "+str(paragraph_count))
print("table count = "+str(table_count))
print(table_matrix_dict)

document.save('new-message.docx')
